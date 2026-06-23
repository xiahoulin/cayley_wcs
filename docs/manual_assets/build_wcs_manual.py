# -*- coding: utf-8 -*-
"""生成《礁盘工业WCS使用功能说明书.docx》。
参考 cayley_wms/docs 的同名说明书结构（python-docx + PIL 渲染图），
按 cayley_wcs 实际代码梳理模块。运行：python build_wcs_manual.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

DOCS = Path(__file__).resolve().parents[1]
RENDER = DOCS / "manual_assets" / "rendered"
RENDER.mkdir(parents=True, exist_ok=True)
OUTPUT = DOCS / "礁盘工业WCS使用功能说明书.docx"

BODY_FONT = "Microsoft YaHei"
TITLE_COLOR = RGBColor(11, 37, 69)
H1_COLOR = RGBColor(46, 116, 181)
H2_COLOR = RGBColor(31, 77, 120)
MUTED = RGBColor(75, 85, 99)

# 调色板（PIL，工业风：灰白底/橙强调/蓝）
C_NAVY = "#0b2545"
C_BLUE = "#2e74b5"
C_BLUE_L = "#e8f1fb"
C_ORANGE = "#f57c00"
C_ORANGE_L = "#fdeede"
C_GRAY = "#6b7280"
C_BORDER = "#c9d4e0"
C_GREEN = "#1f9d55"
C_GREEN_L = "#e7f4ec"
C_WHITE = "#ffffff"
C_BG = "#f4f6f9"


# ---------------- PIL 渲染 ----------------
def font(size, bold=False):
    for path in (("C:/Windows/Fonts/msyhbd.ttc",) if bold else ()) + (
        "C:/Windows/Fonts/msyh.ttc", "C:/Windows/Fonts/simhei.ttf"):
        try:
            return ImageFont.truetype(path, size)
        except OSError:
            continue
    return ImageFont.load_default()


def ctext(d, cx, cy, text, f, fill=C_NAVY):
    bb = d.textbbox((0, 0), text, font=f)
    d.text((cx - (bb[2] - bb[0]) / 2, cy - (bb[3] - bb[1]) / 2), text, font=f, fill=fill)


def ltext(d, x, y, text, f, fill=C_NAVY):
    d.text((x, y), text, font=f, fill=fill)


def box(d, x, y, w, h, title, lines=None, fill=C_WHITE, border=C_BLUE, tcolor=C_NAVY, r=10):
    d.rounded_rectangle([x, y, x + w, y + h], radius=r, fill=fill, outline=border, width=2)
    ft = font(17, bold=True)
    if lines:
        ctext(d, x + w / 2, y + 20, title, ft, tcolor)
        fl = font(13)
        yy = y + 42
        for ln in lines:
            ctext(d, x + w / 2, yy, ln, fl, C_GRAY)
            yy += 19
    else:
        ctext(d, x + w / 2, y + h / 2, title, ft, tcolor)


def arrow(d, x1, y1, x2, y2, color=C_GRAY, width=2):
    d.line([x1, y1, x2, y2], fill=color, width=width)
    import math
    ang = math.atan2(y2 - y1, x2 - x1)
    s = 9
    d.polygon([(x2, y2),
               (x2 - s * math.cos(ang - 0.5), y2 - s * math.sin(ang - 0.5)),
               (x2 - s * math.cos(ang + 0.5), y2 - s * math.sin(ang + 0.5))], fill=color)


def render_business_logic():
    W, H = 1640, 1180
    img = Image.new("RGB", (W, H), C_BG)
    d = ImageDraw.Draw(img)
    ctext(d, W / 2, 34, "CayleyWCS 系统业务逻辑图", font(30, bold=True), C_NAVY)
    ctext(d, W / 2, 66, "上位 WMS  ⇄  WCS 多协议通讯中台  ⇄  下位设备 PLC", font(15), C_GRAY)

    # 三层泳道背景
    lanes = [(100, "上位系统层", C_ORANGE_L, C_ORANGE),
             (360, "WCS 中台（本系统）", C_BLUE_L, C_BLUE),
             (980, "下位设备层", C_GREEN_L, C_GREEN)]
    for y, name, bg, bd in lanes:
        hh = 560 if "中台" in name else 150
        d.rounded_rectangle([40, y, W - 40, y + hh], radius=14, fill=bg, outline=bd, width=2)
        d.text((54, y + 8), name, font=font(14, bold=True), fill=bd)

    # 上位
    box(d, 120, 140, 360, 86, "WMS / MES / ERP", ["仓库管理系统、上位业务"], C_WHITE, C_ORANGE)
    box(d, 1060, 140, 440, 86, "管理前端 (Vue)", ["协议/应用/连接监控/实时看板"], C_WHITE, C_ORANGE)

    # WCS 中台内部
    box(d, 120, 400, 1400, 70, "接入层　Controller (@PostMapping JSON, ApiResponse)　│　管理端 JWT 鉴权　│　开放接口 AppKey 签名 (/open/**)",
        fill=C_WHITE, border=C_BLUE, tcolor=C_NAVY)
    mid = [
        (120, "配置中台", ["协议·点位", "应用(AppKey)", "绑定授权", "数据字典", "故障码"]),
        (470, "任务调度", ["三段握手状态机", "任务引擎(单飞)", "/open/task/dispatch"]),
        (820, "连接治理", ["线程池+信号量", "60s 超时回收", "看门狗重连"]),
        (1170, "监控·报警·审计", ["WS 实时推送", "故障码→报警", "连接/报文日志"]),
    ]
    for x, t, ls in mid:
        box(d, x, 500, 330, 150, t, ls, C_WHITE, C_BLUE, C_NAVY)
    box(d, 120, 678, 660, 70, "适配器工厂 + SPI（Adapter / Factory / Template / Strategy）",
        ["OPC UA · Modbus TCP · S7 (PLC4X)  |  MQTT (Paho)  |  HTTP  |  TCP  |  sim"], C_WHITE, C_BLUE, C_NAVY)
    box(d, 820, 678, 330, 70, "持久化", ["PostgreSQL (jsonb)"], C_WHITE, C_BLUE, C_NAVY)
    box(d, 1190, 678, 330, 70, "缓存", ["Redis (防重放/分布式锁预留)"], C_WHITE, C_BLUE, C_NAVY)

    # 下位设备
    box(d, 150, 1020, 380, 86, "堆垛机 PLC", ["OPC UA  (DB100 三段握手)"], C_WHITE, C_GREEN)
    box(d, 630, 1020, 380, 86, "输送线 / IO", ["Modbus TCP (holding-register…)"], C_WHITE, C_GREEN)
    box(d, 1110, 1020, 380, 86, "提升机 / 其他", ["S7 / MQTT / HTTP"], C_WHITE, C_GREEN)

    # 箭头：上位↔WCS
    arrow(d, 300, 226, 300, 398, C_ORANGE, 3)
    ltext(d, 312, 290, "① 下发任务", font(13, bold=True), C_ORANGE)
    ltext(d, 312, 310, "REST /open/** (AppKey)", font(12), C_GRAY)
    arrow(d, 360, 398, 360, 226, C_ORANGE, 3)
    ltext(d, 372, 340, "② WS 推送 + REST 对账", font(12), C_GRAY)
    arrow(d, 1280, 226, 1280, 398, C_ORANGE, 3)
    ltext(d, 1292, 300, "JWT 管理", font(12), C_GRAY)
    # 适配器↔设备
    for sx, dx, lab in [(300, 340, "OPC UA"), (790, 820, "Modbus"), (1300, 1300, "S7/…")]:
        arrow(d, sx, 748, dx, 1018, C_GREEN, 3)
    ltext(d, 250, 880, "③ 轮询读状态 / 写命令（归一化为统一 JSON 快照）", font(13, bold=True), C_GREEN)

    p = RENDER / "wcs-business-logic.png"
    img.save(p)
    return p


def render_handshake():
    W, H = 1280, 760
    img = Image.new("RGB", (W, H), C_WHITE)
    d = ImageDraw.Draw(img)
    ctext(d, W / 2, 30, "堆垛机取放货 · 三段握手时序图（6.18 协议）", font(24, bold=True), C_NAVY)
    actors = [("WMS / 调度", 200, C_ORANGE), ("WCS 任务引擎", 640, C_BLUE), ("堆垛机 PLC", 1080, C_GREEN)]
    for name, x, c in actors:
        d.rounded_rectangle([x - 130, 70, x + 130, 112], radius=8, fill=C_WHITE, outline=c, width=2)
        ctext(d, x, 91, name, font(15, bold=True), c)
        d.line([x, 112, x, 720], fill=C_BORDER, width=2)

    def msg(y, x1, x2, text, c=C_GRAY, dash=False):
        arrow(d, x1, y, x2, y, c, 2)
        mx = (x1 + x2) / 2
        ctext(d, mx, y - 14, text, font(13), C_NAVY)

    steps = [
        (160, 200, 640, "下发任务 /open/task/dispatch", C_ORANGE),
        (210, 640, 1080, "检查：读 模式=联机自动(2)/故障码=0/任务=0", C_BLUE),
        (260, 640, 1080, "下发参数：排/列/层/口/任务号/类型", C_BLUE),
        (310, 640, 1080, "写 cmd_ConfirmTask=1 (执行任务)", C_BLUE),
        (360, 1080, 640, "回读 status_Task=1 (执行中)", C_GREEN),
        (410, 640, 1080, "写 cmd_ConfirmTask=0 (无意义)", C_BLUE),
        (470, 1080, 640, "回读 status_Task=2 (完成)", C_GREEN),
        (520, 640, 1080, "写 cmd_ConfirmTask=2 (完成确认)", C_BLUE),
        (570, 640, 1080, "清零命令区 (排/列/层/口/号/类型=0)", C_BLUE),
        (630, 640, 200, "WS 推送完成 + /open/task/query 对账", C_ORANGE),
    ]
    for y, x1, x2, t, c in steps:
        msg(y, x1, x2, t, c)
    d.rounded_rectangle([590, 330, 690, 500], radius=6, outline=C_BLUE, width=2)
    ctext(d, 640, 315, "执行中循环", font(12), C_BLUE)
    ctext(d, W / 2, 700, "执行期间读到 status_ErrorCode≠0 → 任务判失败 + 故障码联动报警", font(13), C_ORANGE)
    p = RENDER / "wcs-handshake.png"
    img.save(p)
    return p


def render_state_machine():
    W, H = 1320, 460
    img = Image.new("RGB", (W, H), C_WHITE)
    d = ImageDraw.Draw(img)
    ctext(d, W / 2, 28, "连接治理 · 状态机（需求 4）", font(24, bold=True), C_NAVY)
    states = {
        "NEW": (70, 200), "CONNECTING": (270, 200), "CONNECTED": (500, 130),
        "RUNNING": (740, 130), "DISCONNECTED": (740, 300), "RECONNECTING": (500, 300),
        "FAILED": (1010, 300), "CLOSED": (1010, 130),
    }
    for name, (x, y) in states.items():
        c = C_GREEN if name in ("CONNECTED", "RUNNING") else C_ORANGE if name in ("FAILED", "CLOSED") else C_BLUE
        d.rounded_rectangle([x, y, x + 170, y + 56], radius=10, fill=C_WHITE, outline=c, width=2)
        ctext(d, x + 85, y + 28, name, font(15, bold=True), c)

    def edge(a, b, lab="", c=C_GRAY):
        (x1, y1), (x2, y2) = states[a], states[b]
        arrow(d, x1 + 170 if x2 > x1 else x1, y1 + 28, x2 if x2 > x1 else x2 + 170, y2 + 28, c, 2)
        if lab:
            ctext(d, (x1 + x2) / 2 + 85, (y1 + y2) / 2 + 12, lab, font(12), C_NAVY)

    edge("NEW", "CONNECTING", "open(占槽)")
    edge("CONNECTING", "CONNECTED", "建连OK")
    edge("CONNECTED", "RUNNING", "启动轮询")
    arrow(d, 825, 186, 825, 298, C_GRAY, 2); ctext(d, 900, 240, "心跳失活", font(12), C_NAVY)
    arrow(d, 740, 328, 670, 328, C_GRAY, 2); ctext(d, 600, 270, "看门狗重连", font(12), C_NAVY)
    arrow(d, 585, 300, 700, 160, C_GREEN, 2)
    arrow(d, 910, 328, 1010, 328, C_ORANGE, 2); ctext(d, 960, 360, "重试耗尽", font(12), C_NAVY)
    arrow(d, 355, 254, 460, 310, C_ORANGE, 2); ctext(d, 360, 300, "60s超时/失败→回收槽", font(12), C_ORANGE)
    ctext(d, W / 2, 430, "建连超时(默认60s)取消并释放连接槽；满槽时新建连返回“连接已满，请等待”", font(13), C_GRAY)
    p = RENDER / "wcs-state-machine.png"
    img.save(p)
    return p


# ---------------- docx ----------------
def set_run(run, size=10, bold=False, color=None):
    run.font.name = BODY_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def shade(cell, color):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tcpr.append(shd)


def para(doc, text="", size=10, bold=False, color=None, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    set_run(p.add_run(text), size=size, bold=bold, color=color)
    return p


def heading(doc, text, level=1):
    p = doc.add_heading("", level=level)
    set_run(p.add_run(text), size=16 if level == 1 else 13 if level == 2 else 11,
            bold=True, color=H1_COLOR if level == 1 else H2_COLOR)
    return p


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        shade(c, "E8EEF5")
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(h), size=9, bold=True, color=TITLE_COLOR)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_run(cells[i].paragraphs[0].add_run(str(v)), size=8.5)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def image(doc, path, caption, width=6.4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(cap.add_run(caption), size=8.5, color=MUTED)


# ---------------- 内容（按 cayley_wcs 代码梳理）----------------
MODULES = [
    ("登录与鉴权", "已完成", "管理端 JWT 登录(admin/1)、刷新令牌；开放接口 AppKey 签名+防重放；未认证返回 401"),
    ("协议管理", "已完成", "对接协议(系统/类型字典/JSON 数据格式) + 协议点位/字段映射"),
    ("应用管理", "已完成", "可对接应用/设备，AppKey 生成/重置，绑定协议与连接参数"),
    ("绑定授权", "已完成", "应用绑定授权(wcs_app_binding，上位侧→下位侧；不锁 direction、禁自绑定，按应用整批勾选)；dispatch 越权防护(FORBIDDEN)"),
    ("数据字典", "已完成", "协议类型/方向/故障级别/任务类型/连接状态等枚举单一事实源"),
    ("故障码管理", "已完成", "按协议维护故障码，运行时解析(命中/统一兜底)"),
    ("连接治理", "已完成", "线程池+信号量(最大连接/60s 超时回收)+看门狗重连+状态机+手动读写"),
    ("协议适配", "已完成", "适配器 SPI；OPC UA/Modbus/S7(PLC4X)、MQTT、HTTP、TCP、sim"),
    ("任务调度", "已完成", "WMS 下发(幂等) + 堆垛机三段握手状态机 + 引擎(设备级单飞)"),
    ("监控与报警", "已完成", "WebSocket 实时推送 + 报警生命周期 + 故障码→报警联动"),
    ("审计", "已完成", "连接事件落库 + 报文留存"),
    ("WMS 对账闭环", "已完成", "WS 提示(seq) + REST 增量对账 + snapshot 重同步"),
    ("仿真器", "已完成", "内存 sim + Milo OPC UA 仿真服务端(真·OPC UA 端到端)，无硬件联调 + 故障注入"),
]

# (模块, 功能说明, [接口], [数据表], [关键类])
MODULE_DETAIL = [
    ("登录与鉴权",
     "管理端采用轻量 HS256 JWT：/login 用 admin/1 换取 access_token（默认 120 分钟），/refresh-token 刷新；令牌仅放 Authorization: Bearer。"
     "开放接口 /open/** 面向 WMS/上位系统，使用 AppKey + HMAC-SHA256 签名（X-App-Key/X-Timestamp/X-Nonce/X-Sign）+ Redis 防重放。"
     "未认证统一返回 401（带 ApiResponse JSON）。预留 APISIX/Keycloak 接入开关。",
     ["POST /login", "POST /refresh-token"],
     ["wcs_user", "wcs_userrole"],
     ["JwtService", "JwtAuthenticationFilter", "AppKeyAuthFilter", "AppKeySigner"]),
    ("协议管理",
     "维护“对接协议”：对接系统、协议类型(协议类型字典：opcua/modbus_tcp/s7/tcp/mqtt/http/sim)、协议数据格式(JSON)。"
     "其下维护“协议点位/字段映射”——把设备字段映射到地址(OPC UA NodeId、S7 DB100.DBW0、Modbus holding-register 等)、数据类型、读写、范围。"
     "协议类型驱动适配器选择；点位表是“协议数据格式”的真正落地处。",
     ["POST /protocol/list | /all | /detail | /create | /update | /delete",
      "POST /protocol/point/list | /create | /update | /delete"],
     ["wcs_protocol", "wcs_protocol_point"],
     ["ProtocolController", "ProtocolService"]),
    ("应用管理",
     "维护可对接应用/设备(立体库、输送线…)。每个应用自动生成 AppKey/AppSecret，绑定一个协议(protocol_id)，配置方向(上位/下位)与连接参数 conn_params(host/port/endpoint/unitId 等 JSON)。"
     "出站建立连接前先校验 AppKey 有效且启用(开闸)；可重置 AppSecret。",
     ["POST /application/list | /all | /detail | /create | /update | /delete | /reset-secret"],
     ["wcs_application"],
     ["ApplicationController", "ApplicationService"]),
    ("绑定授权",
     "防越权：阻止“已持有某 key 的应用借 WCS 去指挥本无权指挥的别的应用/设备”。维护“上位侧应用→下位侧应用”的指挥授权表 wcs_app_binding。"
     "/open/task/dispatch 经 AppKey 验签得到调用方真实身份(ATTR_APP_ID，证明)，与请求体 appId(声明) 一起查绑定：仅当存在启用中的绑定才放行，否则 FORBIDDEN。"
     "模型：不按应用 direction 锁方向——任意应用都可出现在上位侧或下位侧(含“下位→下位”)，仅禁止自绑定(upstream==downstream)；direction 仅作提示。"
     "推荐用法——按应用授权：种子 WMS-MAIN(ak_wms_main) 用自身 AppKey 调 /open/**；在“绑定授权”页选定一个上位侧应用，勾选它可指挥的下位侧应用(多选，已排除自身)，保存即整批授权(grant 声明式对账：缺补、撤选软删、自指忽略、幂等)。",
     ["POST /binding/list | /all | /detail | /create | /update | /delete",
      "POST /binding/granted（查某上位侧已授权应用）| /binding/grant（按应用整批授权）"],
     ["wcs_app_binding", "wcs_application(WMS-MAIN 上位侧示例)"],
     ["BindingController", "BindingService(grant/grantedDownstreamIds)", "OpenTaskController(assertAllowed)"]),
    ("数据字典",
     "通用数据字典，作为协议类型、对接方向、故障级别、报警状态、任务类型、任务状态、连接状态、点位读写/分组等枚举的单一事实源；前端下拉直接消费。",
     ["POST /dict/type/list | /create | /update | /delete",
      "POST /dict/item/list-by-type | /create | /update | /delete"],
     ["wcs_dict_type", "wcs_dict_item"],
     ["DictController", "DictService"]),
    ("故障码管理",
     "按协议维护故障码枚举(码/级别/名称/信息/处置建议)。运行时 FaultCodeResolver 解析：命中返回维护信息，未维护返回统一兜底“未知故障，故障码=X”。"
     "已内置堆垛机标准协议(6.18 修订) 1–48 故障码。",
     ["POST /fault-code/list | /create | /update | /delete | /resolve"],
     ["wcs_fault_code"],
     ["FaultCodeController", "FaultCodeResolver"]),
    ("连接治理",
     "线程池管理设备连接：信号量限制最大连接数(默认 32，满则返回“连接已满，请等待”)；每次建连有超时(默认 60s)，超时未连成功则取消并回收连接槽；"
     "成功后每连接一条工作线程做心跳+状态轮询；心跳看门狗检测失活并有界自动重连；连接状态机(NEW→CONNECTING→CONNECTED→RUNNING→DISCONNECTED→RECONNECTING→FAILED/CLOSED)。"
     "提供手动建连/断开/重连、单点读写、连接槽与状态查询。",
     ["POST /connection/open | /close | /reconnect | /detail | /status | /read | /write"],
     ["—（运行态，不落业务表）"],
     ["ConnectionManager", "ManagedConnection", "ConnectionWatchdog"]),
    ("协议适配",
     "统一适配器 SPI：到连接阶段，所有协议都用同一套接口处理“连接 + 读写”，读出归一化为统一 JSON 快照（field_name→value）。"
     "设计模式：Adapter(屏蔽差异) + Factory(按 protocol_type 选 Provider) + Template Method(固化生命周期) + Strategy(JSON 编解码 DataCodec)。"
     "实现：OPC UA / Modbus TCP / S7 由 Apache PLC4X 统一（连接串由 conn_params.endpoint 拼装，点位 address 直接当 PLC4X tag/NodeId）；MQTT 用 Eclipse Paho；HTTP 用 JDK HttpClient；TCP 用 Socket；loopback/sim 为内存仿真。"
     "OPC UA 由 StackerOpcUaAdapter（继承 Plc4xAdapter）实现协议特有“主动心跳”：每 5s 翻转写 WCS_Heart(0↔1) 给 PLC；心跳点也从点位表按 field_name 定位、不写死。"
     "PLC4X 的 OPC UA 通道非并发安全，故 read/write/连接在 channelLock 上串行化（轮询线程与手动/握手读写共用一条会话）。",
     ["—（被连接治理调用，无独立 HTTP 接口）"],
     ["—"],
     ["ProtocolAdapter / ProtocolAdapterFactory / ProtocolAdapterProvider", "Plc4xAdapter（OPC UA/Modbus/S7 共用，channelLock 串行）",
      "StackerOpcUaAdapter（OPC UA 主动心跳）/ OpcUaAdapterProvider / Plc4xConnectionStrings", "MqttAdapter / HttpAdapter / TcpAdapter / SimAdapter / DataCodec"]),
    ("任务调度",
     "接收 WMS 经 /open/task/dispatch 下发的任务(幂等：同 app+task_no 不重复；并校验调用方对目标设备的绑定授权，越权返回 FORBIDDEN)。任务引擎按设备级单飞推进堆垛机三段握手状态机："
     "检查(模式=联机自动/无故障/无任务) → 下发参数(排/列/层/口/任务号/类型) → 写执行确认 → 等待完成 → 写完成确认 → 清零命令区。"
     "执行期间读到故障码即判失败并联动报警。支持任务 CRUD 与取消。",
     ["POST /task/list | /detail | /create | /cancel", "POST /open/task/dispatch (AppKey)"],
     ["wcs_task"],
     ["TaskEngine", "StackerHandshakeStateMachine", "TaskService"]),
    ("监控与报警",
     "WebSocket /ws/monitor 每秒推送一帧(带 seq+ts)：所有连接的实时态(心跳/模式/任务/速度/位置/故障码)+连接槽+活动报警。"
     "报警生命周期：产生/确认/清除/历史 + 去重。故障码→报警联动：轮询检测 status_ErrorCode 边沿 → 查故障码表 → 产生报警(未维护走兜底)；连接断开/恢复联动通讯报警。",
     ["POST /alarm/list | /active | /ack | /clear", "WS /ws/monitor"],
     ["wcs_alarm"],
     ["MonitorPusher", "MonitorWebSocketHandler", "AlarmService", "AlarmEventListener"]),
    ("审计",
     "连接事件审计：建连/断开/超时/重连等状态变化落库 wcs_connection_log。报文留存：手动/任务读写的点位值记入 wcs_message_log（高频，按需保留）。",
     ["POST /audit/connection-log/list | /audit/message-log/list"],
     ["wcs_connection_log", "wcs_message_log"],
     ["AuditService", "ConnectionAuditListener"]),
    ("WMS 对账闭环",
     "WMS↔WCS 采用“命令走 REST、遥测走 WS、关键事件 REST 对账”。WS 帧带单调 seq，断号即触发对账；"
     "/open/task/query、/open/alarm/query 按 last_update_time 水位线增量拉变化；/open/snapshot 断线后一把全量重同步。保证任务完成/报警不丢(最终一致)。",
     ["POST /open/task/query | /open/alarm/query | /open/snapshot (AppKey)"],
     ["wcs_task", "wcs_alarm"],
     ["OpenReconcileController"]),
    ("仿真器",
     "两种无硬件联调手段："
     "①【内存仿真】协议类型 sim 的 SimAdapter + StackerDeviceState：进程内假读写，跑通连接治理/三段握手/报警，不开端口。"
     "②【Milo OPC UA 仿真服务端】OpcUaStackerSimulator（Spring profile simulator）：后端进程内内嵌一个真正的 Eclipse Milo OPC UA Server，"
     "监听 opc.tcp://127.0.0.1:4840/cayleywcs/stacker（ns=2），按 DB100 字段注册 UaVariableNode，背后挂内存设备状态机每秒 tick（心跳翻转/握手应答/故障注入）。"
     "WCS 经 PLC4X 当真·OPC UA 客户端连它 → 真·OPC UA 端到端联调（连接/心跳/读状态/写命令/三段握手/注故障）。"
     "对应种子：协议 PALLET_SIM_OPCUA + 应用 PALLET_SIM(V5)；注故障 POST /simulator/inject-fault {appId:999001}。"
     "两种仿真都支持注入/清除故障码、复位、查看内部态。",
     ["POST /simulator/inject-fault | /clear-fault | /reset | /snapshot"],
     ["—（内存态 / OPC UA 服务端）"],
     ["OpcUaStackerSimulator（Milo OPC UA Server, profile simulator）", "StackerNamespace（NodeId 注册）",
      "StackerDeviceState / StackerSimulatorRegistry", "SimAdapter（内存 sim 适配器）"]),
]

INTERFACES = [
    ("登录与鉴权", "/login", "管理员登录获取 JWT", "公开"),
    ("登录与鉴权", "/refresh-token", "刷新访问令牌", "公开"),
    ("数据字典", "/dict/type/list /create /update /delete", "字典类型 CRUD", "JWT"),
    ("数据字典", "/dict/item/list-by-type /create /update /delete", "字典项 CRUD", "JWT"),
    ("协议管理", "/protocol/list /all /detail /create /update /delete", "协议 CRUD", "JWT"),
    ("协议管理", "/protocol/point/list /create /update /delete", "协议点位 CRUD", "JWT"),
    ("应用管理", "/application/list /all /detail /create /update /delete /reset-secret", "应用与 AppKey", "JWT"),
    ("绑定授权", "/binding/list /all /detail /create /update /delete", "上下位绑定授权 CRUD", "JWT"),
    ("绑定授权", "/binding/granted /binding/grant", "按应用查/整批授权（上位侧→下位侧）", "JWT"),
    ("故障码管理", "/fault-code/list /create /update /delete /resolve", "故障码 CRUD + 解析", "JWT"),
    ("连接治理", "/connection/open /close /reconnect /detail /status /read /write", "建连/断开/重连/读写/状态", "JWT"),
    ("任务调度", "/task/list /detail /create /cancel", "任务 CRUD/取消", "JWT"),
    ("任务调度", "/open/task/dispatch", "WMS 下发任务(幂等 + 绑定越权校验)", "AppKey"),
    ("监控与报警", "/alarm/list /active /ack /clear", "报警列表/确认/清除", "JWT"),
    ("监控与报警", "/ws/monitor", "实时状态/报警推送(带 seq)", "公开(内网)"),
    ("审计", "/audit/connection-log/list /audit/message-log/list", "连接/报文日志", "JWT"),
    ("WMS 对账", "/open/task/query /open/alarm/query /open/snapshot", "增量对账 + 重同步", "AppKey"),
    ("仿真器", "/simulator/inject-fault /clear-fault /reset /snapshot", "仿真注故障/复位/查看", "JWT"),
    ("健康", "/health", "健康检查", "公开"),
]

TABLES = [
    ("认证", "wcs_user / wcs_userrole", "管理员账号与角色"),
    ("配置中台", "wcs_dict_type / wcs_dict_item", "数据字典(协议类型等枚举)"),
    ("配置中台", "wcs_protocol / wcs_protocol_point", "协议表 + 点位/字段映射(协议数据格式)"),
    ("配置中台", "wcs_application", "应用信息表(AppKey、绑定协议、连接参数)"),
    ("配置中台", "wcs_app_binding", "应用绑定授权(上位侧→下位侧，不锁 direction、禁自绑定，下发越权防护)"),
    ("配置中台", "wcs_fault_code", "按协议的故障码枚举"),
    ("任务调度", "wcs_task", "WCS 任务(下发 + 握手执行态)"),
    ("监控报警", "wcs_alarm", "报警(产生/确认/清除/历史)"),
    ("审计", "wcs_connection_log / wcs_message_log", "连接事件 + 报文留存"),
]

# ===== 全部表完整字段结构（列名 / 类型 / 说明），源自 Flyway V1 + V6 =====
_STD = [
    ("creator", "varchar(512)", "创建者"),
    ("create_time", "timestamp(6)", "创建时间"),
    ("last_update_time", "timestamp(6)", "最后更新时间"),
    ("is_valid", "boolean", "软删标志（true=有效）"),
    ("tenant_id", "bigint", "租户 id（默认 1）"),
]
_ID = ("id", "bigint", "主键，自增")

# (表名, 用途, [(列, 类型, 说明)])
TABLE_SCHEMAS = [
    ("wcs_userrole", "角色（登录用）", [
        _ID,
        ("role_name", "varchar(512)", "角色名"),
        ("is_valid", "boolean", "软删标志"),
        ("create_time", "timestamp(6)", "创建时间"),
        ("last_update_time", "timestamp(6)", "最后更新时间"),
        ("tenant_id", "bigint", "租户 id"),
    ]),
    ("wcs_user", "管理员账号（JWT 登录）", [
        _ID,
        ("user_num", "varchar(512)", "登录账号（如 admin）"),
        ("user_name", "varchar(512)", "姓名"),
        ("contact_tel", "varchar(512)", "联系电话"),
        ("user_role", "varchar(512)", "角色"),
        ("sex", "varchar(512)", "性别"),
        ("auth_string", "varchar(512)", "口令哈希"),
        ("email", "varchar(512)", "邮箱"),
        ("is_valid", "boolean", "软删标志"),
        ("creator", "varchar(512)", "创建者"),
        ("create_time", "timestamp(6)", "创建时间"),
        ("last_update_time", "timestamp(6)", "最后更新时间"),
        ("tenant_id", "bigint", "租户 id"),
    ]),
    ("wcs_dict_type", "字典类型", [
        _ID,
        ("type_code", "varchar(128)", "类型编码（如 protocol_type）"),
        ("type_name", "varchar(256)", "类型名称"),
        ("remark", "varchar(512)", "备注"),
        ("sort", "bigint", "排序"),
    ] + _STD),
    ("wcs_dict_item", "字典项（枚举值）", [
        _ID,
        ("type_code", "varchar(128)", "所属字典类型编码"),
        ("item_code", "varchar(128)", "项编码"),
        ("item_name", "varchar(256)", "项名称"),
        ("item_value", "varchar(512)", "项值"),
        ("sort", "bigint", "排序"),
        ("remark", "varchar(512)", "备注"),
    ] + _STD),
    ("wcs_protocol", "协议表（对接协议）", [
        _ID,
        ("protocol_code", "varchar(128)", "协议编码（业务键）"),
        ("protocol_name", "varchar(256)", "协议名称"),
        ("target_system", "varchar(256)", "对接系统"),
        ("protocol_type", "varchar(64)", "协议类型字典码（opcua/modbus_tcp/s7/tcp/mqtt/http/sim）→ 决定适配器"),
        ("data_format", "text(JSON)", "数据格式（块/访问/寻址规则）"),
        ("transport_template", "text(JSON)", "传输模板（如 NodeId 模式）"),
        ("version", "varchar(64)", "版本"),
        ("status", "varchar(32)", "状态（enabled/disabled）"),
        ("remark", "varchar(1024)", "备注"),
    ] + _STD),
    ("wcs_protocol_point", "协议点位/字段映射（协议数据格式落地）", [
        _ID,
        ("protocol_id", "bigint", "所属协议 id"),
        ("point_group", "varchar(64)", "分区（From_WCS/To_WCS）"),
        ("field_name", "varchar(128)", "逻辑字段名（统一 JSON 的 key）"),
        ("symbol_name", "varchar(256)", "符号名"),
        ("address", "varchar(256)", "驱动级地址：OPC UA NodeId / S7 DB100.DBW0 / Modbus 寄存器"),
        ("data_type", "varchar(64)", "数据类型（INT/UINT/REAL/WORD/BOOL…）"),
        ("rw", "varchar(2)", "读写（R/W）"),
        ("value_range", "varchar(256)", "取值范围"),
        ("scale", "numeric(18,6)", "缩放系数（当前仅存，未自动应用）"),
        ("sort", "bigint", "排序"),
        ("description", "varchar(1024)", "描述"),
    ] + _STD),
    ("wcs_application", "应用信息表（可对接应用/设备 + AppKey）", [
        _ID,
        ("app_code", "varchar(128)", "应用编码"),
        ("app_name", "varchar(256)", "应用名称"),
        ("app_key", "varchar(128)", "AppKey（鉴权用，建索引）"),
        ("app_secret", "varchar(256)", "AppSecret（HMAC 签名密钥）"),
        ("protocol_id", "bigint", "绑定协议 id"),
        ("direction", "varchar(32)", "方向（upstream 上位 / downstream 下位）"),
        ("conn_params", "text(JSON)", "连接参数（endpoint/host/port/unitId…）"),
        ("max_retry", "bigint", "最大重连次数"),
        ("heartbeat_interval_ms", "bigint", "心跳/轮询间隔(ms)"),
        ("enabled", "boolean", "是否启用"),
        ("status", "varchar(32)", "状态"),
        ("remark", "varchar(1024)", "备注"),
    ] + _STD),
    ("wcs_app_binding", "应用绑定授权（上位侧→下位侧，越权防护）。唯一索引 (upstream_app_id, downstream_app_id, tenant_id)；禁自绑定", [
        _ID,
        ("upstream_app_id", "bigint", "上位侧应用 id（调用方）"),
        ("downstream_app_id", "bigint", "下位侧应用 id（目标）"),
        ("scope", "varchar(32)", "权限范围（dispatch/read/control）"),
        ("enabled", "boolean", "是否启用"),
        ("remark", "varchar(1024)", "备注"),
    ] + _STD),
    ("wcs_fault_code", "故障码表（按协议维护）", [
        _ID,
        ("protocol_id", "bigint", "所属协议 id"),
        ("code", "bigint", "故障码"),
        ("level", "varchar(32)", "级别（error/warn…）"),
        ("name", "varchar(256)", "名称"),
        ("message", "varchar(1024)", "故障信息"),
        ("suggestion", "varchar(2048)", "处置建议"),
    ] + _STD),
    ("wcs_task", "WCS 任务（下发 + 握手执行态）", [
        _ID,
        ("task_no", "varchar(128)", "任务号（与 app 组合幂等去重）"),
        ("app_id", "bigint", "目标设备应用 id"),
        ("wms_ref", "varchar(256)", "WMS 单号/引用"),
        ("task_type", "varchar(64)", "任务类型（入库/出库/移库…字典码）"),
        ("take_row", "bigint", "取货-排"),
        ("take_column", "bigint", "取货-列"),
        ("take_floor", "bigint", "取货-层"),
        ("put_row", "bigint", "放货-排"),
        ("put_column", "bigint", "放货-列"),
        ("put_floor", "bigint", "放货-层"),
        ("port_num", "bigint", "出入口号"),
        ("priority", "bigint", "优先级"),
        ("status", "varchar(32)", "状态（pending/dispatched/executing/completed/cancelled）"),
        ("handshake_step", "varchar(64)", "握手步（CREATED→…→DONE）"),
        ("error_code", "bigint", "故障码（执行失败）"),
        ("payload", "text(JSON)", "附加载荷"),
        ("dispatch_time", "timestamp(6)", "下发时间"),
        ("finish_time", "timestamp(6)", "完成时间"),
    ] + _STD),
    ("wcs_alarm", "报警（产生/确认/清除/历史）", [
        _ID,
        ("app_id", "bigint", "应用/设备 id"),
        ("fault_code", "bigint", "故障码"),
        ("level", "varchar(32)", "级别"),
        ("message", "varchar(1024)", "报警信息"),
        ("suggestion", "varchar(2048)", "处置建议"),
        ("status", "varchar(32)", "状态（raised/ack/cleared）"),
        ("raised_time", "timestamp(6)", "产生时间"),
        ("ack_by", "varchar(256)", "确认人"),
        ("ack_time", "timestamp(6)", "确认时间"),
        ("cleared_time", "timestamp(6)", "清除时间"),
    ] + _STD),
    ("wcs_connection_log", "连接事件审计（无标准列 creator/is_valid/last_update_time）", [
        _ID,
        ("app_id", "bigint", "应用 id"),
        ("app_code", "varchar(128)", "应用编码"),
        ("event", "varchar(64)", "事件（connect/disconnect/timeout/error/reconnect）"),
        ("state", "varchar(64)", "连接状态"),
        ("detail", "varchar(2048)", "明细"),
        ("create_time", "timestamp(6)", "时间"),
        ("tenant_id", "bigint", "租户 id"),
    ]),
    ("wcs_message_log", "报文留存（高频，按需清理；无标准列 creator/is_valid/last_update_time）", [
        _ID,
        ("app_id", "bigint", "应用 id"),
        ("direction", "varchar(16)", "方向（read/write）"),
        ("field_name", "varchar(128)", "字段名"),
        ("address", "varchar(256)", "地址（NodeId/寄存器）"),
        ("raw_value", "varchar(512)", "原始值"),
        ("json_payload", "text(JSON)", "JSON 载荷"),
        ("create_time", "timestamp(6)", "时间"),
        ("tenant_id", "bigint", "租户 id"),
    ]),
]


def build():
    bl = render_business_logic()
    hs = render_handshake()
    sm = render_state_machine()

    doc = Document()
    # 封面
    para(doc, "礁盘工业WCS使用功能说明书", size=24, bold=True, color=TITLE_COLOR, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "仓库控制系统 · 多协议设备对接 / 连接治理 / 任务调度 / 实时监控",
         size=12, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "当前范围：登录鉴权、协议/应用/字典/故障码配置、绑定授权、连接治理、协议适配、任务调度、监控报警、审计、WMS 对账、仿真器",
         size=9, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    table(doc, ["版本", "日期", "作者", "说明"], [
        ("V0.1", "2026-06-19", "夏侯霖", "WCS 初版功能说明：按当前代码梳理全部模块、接口、数据表、业务逻辑图与核心时序/状态机。"),
        ("V0.2", "2026-06-23", "夏侯霖", "新增“绑定授权”模块(wcs_app_binding + /binding/* + dispatch 越权防护)；补第 11 节安全边界；同步迁移清单(V4/V5/V6)与测试数(18)。"),
        ("V0.3", "2026-06-23", "夏侯霖", "绑定授权升级为“以上位为中心授权”：种子上位 WMS-MAIN(V7) + /binding/grant 整批授权 + 前端按上位勾选下位设备；测试数(20)。"),
        ("V0.4", "2026-06-23", "夏侯霖", "对抗式评审加固：create() 复用软删行(避免唯一索引冲突)、grant() 按 tenant 隔离+scope 收敛、appkey 关闭记 WARN、前端校正已删上位选择；测试数(23)。"),
        ("V0.5", "2026-06-23", "夏侯霖", "去除自绑定(V8)：绑定保持上位侧→下位侧称呼但不锁 direction，任意应用皆可任一侧(含下位→下位)，禁自绑定；测试数(24)。"),
        ("V0.6", "2026-06-23", "夏侯霖", "二轮评审加固：update() 改方向校验软删行(干净 CONFLICT)、dispatch 越权判定 isAllowed 按调用方租户隔离(读写一致)；测试数(26)。"),
        ("V0.7", "2026-06-24", "夏侯霖", "第 8 节补全：全部 13 张表的完整字段结构（列名/类型/说明），前端拆「授权绑定/绑定明细」两页。"),
    ])

    heading(doc, "1. 项目概述", 1)
    para(doc, "CayleyWCS（仓库控制系统）定位在 WMS（上位业务）与设备（下位 PLC：堆垛机/输送线/提升机等）之间，"
              "负责多协议设备通讯、连接治理、任务调度与实时监控。架构对齐 CayleyWMS。")
    table(doc, ["项", "取值"], [
        ("后端", "Java 25 · Spring Boot 4.0.6 · MyBatis-Plus · PostgreSQL 17 · Redis · Flyway · springdoc"),
        ("工业驱动", "Apache PLC4X(OPC UA/Modbus/S7) · Eclipse Paho(MQTT) · JDK HttpClient · Socket · Eclipse Milo(OPC UA 仿真服务端) · 内存 sim"),
        ("前端", "Vue 3 · Vite · TypeScript（登录 + 协议/应用/授权绑定/绑定明细/连接监控/实时看板）"),
        ("端口", "后端 20021 · 前端 5184 · PostgreSQL 5433 · Redis 6380（避开 WMS）"),
        ("接口约定", "业务接口全 POST + application/json；统一响应 ApiResponse(isSuccess/code/errorMessage/data)"),
        ("鉴权", "管理端 JWT(Bearer)；开放接口 /open/** 用 AppKey + HMAC 签名 + 防重放"),
    ])

    heading(doc, "2. 系统业务逻辑图", 1)
    para(doc, "下图为 WCS 总体业务逻辑：上位 WMS 通过 REST 下发任务、通过 WebSocket+REST 对账接收遥测；"
              "WCS 中台完成鉴权、配置、任务握手、连接治理、协议适配、监控报警与审计；下位设备按各自协议(OPC UA/Modbus/S7…)被统一纳管。")
    image(doc, bl, "图 2.1：CayleyWCS 系统业务逻辑图", 6.6)

    heading(doc, "3. 模块清单", 1)
    table(doc, ["模块", "状态", "能力概述"], [(m, s, d) for m, s, d in MODULES], widths=[1.3, 0.8, 4.3])

    heading(doc, "4. 各业务模块说明", 1)
    for idx, (name, desc, apis, tbls, classes) in enumerate(MODULE_DETAIL, start=1):
        heading(doc, f"4.{idx} {name}", 2)
        para(doc, desc)
        table(doc, ["项", "内容"], [
            ("关键接口", "\n".join(apis)),
            ("数据表", "、".join(tbls)),
            ("关键类", "、".join(classes)),
        ], widths=[1.1, 5.3])

    heading(doc, "5. 堆垛机任务握手时序图", 1)
    para(doc, "WCS 任务调度的核心：堆垛机取放货三段握手（6.18 协议）。命令/状态均为锁存，PLC 等 WCS 逐步确认，因此握手状态不丢。")
    image(doc, hs, "图 5.1：堆垛机取放货三段握手时序图", 6.6)

    heading(doc, "6. 连接治理状态机", 1)
    para(doc, "每个设备连接是一台被治理的连接：信号量限最大连接数、60s 建连超时回收连接槽、心跳看门狗失活自动重连。")
    image(doc, sm, "图 6.1：连接治理状态机", 6.6)

    heading(doc, "7. 接口清单", 1)
    para(doc, "业务接口一律 POST + application/json；下表按模块汇总（鉴权列：JWT=管理端令牌、AppKey=开放接口签名、公开=免鉴权）。")
    table(doc, ["模块", "接口路径", "说明", "鉴权"], [(m, p, d, a) for m, p, d, a in INTERFACES],
          widths=[1.1, 2.7, 1.9, 0.7])

    heading(doc, "8. 数据库表结构", 1)
    para(doc, "PostgreSQL 17（测试用 H2 PostgreSQL 模式）。Flyway 迁移 V1 建表 / V2 字典+管理员 / V3 堆垛机协议+点位+故障码 / "
              "V4 码垛机 OPC UA 协议+应用 / V5 OPC UA 仿真应用 / V6 应用绑定授权表 / V7 上位 WMS-MAIN+授权 / V8 去除自绑定。"
              "多数业务表含标准列：id、creator、create_time、last_update_time、is_valid(软删)、tenant_id；JSON 字段以 text 存储(JacksonTypeHandler 序列化)。")

    heading(doc, "8.1 表清单总览", 2)
    table(doc, ["业务域", "数据表", "说明"], [(d, t, c) for d, t, c in TABLES], widths=[1.2, 2.6, 2.6])

    heading(doc, "8.2 各表字段结构", 2)
    para(doc, "下列为每张表的完整字段（列名 / 类型 / 说明）。text(JSON) 表示以 text 存储、应用层序列化为 JSON。",
         size=9, color=MUTED)
    for idx, (tname, purpose, cols) in enumerate(TABLE_SCHEMAS, start=1):
        heading(doc, f"8.2.{idx}　{tname}", 3)
        para(doc, purpose, size=9, color=MUTED)
        table(doc, ["列名", "类型", "说明"], [(c, t, d) for c, t, d in cols], widths=[1.8, 1.4, 3.2])

    heading(doc, "9. 运行与部署", 1)
    table(doc, ["场景", "命令/配置", "说明"], [
        ("起依赖", "docker compose up -d postgres redis", "PostgreSQL(5433) + Redis(6380)"),
        ("起后端", 'cd backend && mvn spring-boot:run "-Dspring-boot.run.profiles=local"', "默认 20021"),
        ("无硬件 OPC UA 联调", 'mvn spring-boot:run "-Dspring-boot.run.profiles=local,simulator"', "随后端起 Milo OPC UA 仿真服务端 opc.tcp://127.0.0.1:4840/cayleywcs/stacker；对 PALLET_SIM 建连即真·OPC UA 端到端"),
        ("起前端", "cd frontend && npm install && npm run dev", "5184(被占则换端口)；/api 与 /ws 代理到后端"),
        ("完整 Compose", "docker compose --profile full up -d --build", "pg/redis/backend/frontend 分容器；后端默认带 simulator profile(CAYLEYWCS_BACKEND_PROFILES，对接真机置空)"),
        ("Swagger", "/swagger-ui.html", "接口文档"),
        ("测试", "cd backend && mvn test", "H2，无需外部依赖，26 用例"),
    ])

    heading(doc, "10. WMS 对接与对账闭环", 1)
    para(doc, "下发：WMS 调 /open/task/dispatch(AppKey 签名 + 绑定授权)，WCS 幂等建任务并自动跑三段握手。"
              "接收：WMS 订阅 /ws/monitor 拿低延迟提示(带 seq)，断号/重连后用 /open/snapshot 全量重同步、再按 last_update_time 水位线增量 /open/task|alarm/query，按 id 去重。"
              "如此既享受 WS 实时、又用 REST 保证任务完成/报警不丢。")

    heading(doc, "11. 安全：下发越权防护与设备保护边界", 1)
    para(doc, "WCS 安全分两层：① WCS API 层访问控制(AppKey 鉴权 + 绑定授权)，管住“谁能经 WCS 下发、且只能指挥被授权的设备”；"
              "② 设备级保护(网络/PLC 配置)，管住“谁能直接连到下位机”。前者由本系统实现，后者主战场在网络与 PLC，WCS 通过“做成唯一闸口 + 持证书加密连接 + 审计”参与。")
    para(doc, "① API 层(本系统)：/open/** 必须带合法 AppKey 签名(应用存在且启用 + 时间戳未过期 + nonce 不重放 + 签名正确)方可进入；"
              "/open/task/dispatch 进一步做绑定授权——以验签得到的调用方真实身份(ATTR_APP_ID，证明) + 请求体 appId(声明) 查 wcs_app_binding，仅启用中的绑定放行，否则 FORBIDDEN；判定按调用方租户隔离(与唯一索引含 tenant 对齐，读写一致)。"
              "授权按应用维护：在“绑定授权/绑定明细”页选定一个上位侧应用(如 WMS-MAIN)，勾选其可指挥的下位侧应用(整批 grant)；不锁 direction、禁自绑定。由此“调用方只能指挥被授权的应用”，杜绝混淆代理式越权。")
    para(doc, "② 设备级保护(WCS 之外，须现场落实)：")
    table(doc, ["层", "措施", "归属"], [
        ("网络隔离(最重要)", "PLC 置独立 OT/控制网；防火墙仅放行 WCS→PLC 端口，其余不可达", "网络/现场"),
        ("设备侧认证", "OPC UA 启用 SecurityPolicy(Basic256Sha256)+Sign&Encrypt+证书信任，或用户名口令", "PLC 配置"),
        ("Modbus/S7", "协议本身无认证，仅靠网络隔离 + 网关兜底", "网络"),
        ("设备写保护", "PLC 块只读、RUN/PROG 硬件钥匙开关(物理禁写)", "设备"),
        ("硬线安全", "急停/互锁/光幕，不在通讯链路上；WCS 不承担安全停机", "电气安全"),
    ], widths=[1.4, 3.6, 1.2])
    para(doc, "边界提醒：绑定授权属 API 层逻辑访问控制，拦截“经 WCS 越权指挥别人设备”，但拦不住绕开 WCS 直连 PLC 端口者——后者必须靠网络分段把 WCS 做成唯一通道。"
              "此外当前 AppKey 签名未覆盖请求体，坐标/任务类型理论可被篡改，彻底闭合需把 body 摘要并入签名(列为后续)。",
         color=MUTED)

    doc.core_properties.title = "礁盘工业WCS使用功能说明书"
    doc.core_properties.author = "夏侯霖"
    doc.core_properties.subject = "CayleyWCS 功能、接口、数据表、业务逻辑图与时序/状态机说明"
    doc.save(OUTPUT)
    print("written:", OUTPUT)


if __name__ == "__main__":
    build()
